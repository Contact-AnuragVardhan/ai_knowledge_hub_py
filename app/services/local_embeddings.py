# app/services/local_embeddings.py

import time
from typing import List

from fastapi import UploadFile
from pypdf import PdfReader
from docx import Document
from sentence_transformers import SentenceTransformer

from .vector_store import VectorStore
from app.config import settings
from app.utils.logging import logger

# 768-dim, very stable and accurate
_EMBEDDING_MODEL_NAME = "BAAI/bge-base-en-v1.5"

MAX_DOC_SIZE = 5_000_000

logger.info(f"Loading local embedding model: {_EMBEDDING_MODEL_NAME}")
_embedding_model = SentenceTransformer(_EMBEDDING_MODEL_NAME)
logger.info("Local embedding model loaded successfully")


def extract_text(file: UploadFile) -> str:
    filename = file.filename or "unknown"
    lower = filename.lower()
    logger.info(f"Extracting text from file: {filename}")

    try:
        if lower.endswith(".pdf"):
            logger.debug("Detected PDF file type")
            reader = PdfReader(file.file)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)

        elif lower.endswith(".docx"):
            logger.debug("Detected DOCX file type")
            doc = Document(file.file)
            text = "\n".join(p.text for p in doc.paragraphs)

        else:
            logger.debug("Detected text/other file type")
            text = file.file.read().decode("utf-8")

        logger.info(f"Extracted {len(text)} characters from {filename}")
        return text
    except Exception as exc:
        logger.exception(f"Failed to extract text from {filename}: {exc}")
        raise

def extract_text_from_path(path: str, filename: str) -> str:
    """
    Same logic as extract_text, but works on a stored file path.
    Used by background ingest jobs.
    """
    logger.info(f"Extracting text from path={path}, filename={filename}")
    lower = filename.lower()

    try:
        if lower.endswith(".pdf"):
            logger.debug("Detected PDF file type (path)")
            reader = PdfReader(path)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)

        elif lower.endswith(".docx"):
            logger.debug("Detected DOCX file type (path)")
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs)

        else:
            logger.debug("Detected text/other file type (path)")
            with open(path, "rb") as f:
                text = f.read().decode("utf-8", errors="ignore")

        logger.info(f"Extracted {len(text)} characters from {filename} (path)")
        return text
    except Exception as exc:
        logger.exception(f"Failed to extract text from path {path}: {exc}")
        raise

def embed_text(content: str) -> List[float]:
    """
    Create a 768-dim embedding using local SentenceTransformer model.
    This **replaces** the old OpenAI text-embedding-3-small call.
    """
    try:
        logger.info(f"Creating local embedding for content length={len(content)}")

        if not content:
            logger.warning("embed_text called with empty content")
            return []

        # Returns a numpy array of shape (768,)
        vec = _embedding_model.encode(
            content,
            convert_to_numpy=True,
            normalize_embeddings=False,
        )

        embedding = vec.astype(float).tolist()
        logger.debug(f"Local embedding created successfully; dim={len(embedding)}")
        return embedding
    except Exception as exc:
        logger.exception(f"Local embedding creation failed: {exc}")
        raise

def chunk_and_store(user_id: int, doc_name: str, text: str, store: VectorStore) -> None:
    """
    Split the text into chunks, embed each chunk locally, and store in Postgres.
    Assumes document_chunks.embedding is vector(768) in the DB.
    """
    logger.info(
        f"Chunking and storing document: user_id={user_id}, doc_name={doc_name}, "
        f"text_length={len(text)}"
    )

    # Safety cap for very large docs
    text = text[:MAX_DOC_SIZE]

    chunks = [
        text[i:i + settings.chunk_size].strip()
        for i in range(0, len(text), settings.chunk_size)
        if text[i:i + settings.chunk_size].strip()
    ]

    logger.info(f"Total chunks to store for '{doc_name}': {len(chunks)}")

    for idx, chunk in enumerate(chunks):
        #logger.debug(f"Embedding chunk {idx + 1}/{len(chunks)} for '{doc_name}'")
        logger.info(f"Embedding chunk {idx + 1}/{len(chunks)} for '{doc_name}'")

        vec = embed_text(chunk)
        #logger.debug(f"Chunk {idx} embedding dim={len(vec)}")
        logger.info(f"Chunk {idx} embedding dim={len(vec)}")

        store.insert_chunk(user_id, doc_name, idx, chunk, vec)

    store.db.commit()
    logger.info(f"Completed chunking and storing for '{doc_name}'")
